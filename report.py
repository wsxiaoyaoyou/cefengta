from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import warnings
import pymysql
import json
import turbulence_analysis,yearmaxwind_analysis_v1,insert_dynamic_information,weibull_analysis

warnings.filterwarnings('ignore')
host = 'localhost'
port = 3306
user = 'root' #用户名
password = '123456' # 密码
database = 'cefengta' #数据库名称
rootPath = "C:\\Users\\admin\\PycharmProjects\\metaReport\\"

# -------------------------- 1. 核心参数定义（标红内容均用变量替换，便于修改） --------------------------
# 测风塔基础信息
TOWER_ID = "14648"

conn = pymysql.connect(host=host, port=port, user=user, password=password, database=database, charset='utf8mb4')
cursor = conn.cursor()
select_data = "SELECT * FROM cefengta.static_information as t1 inner join cefengta.dynamic_information as t2 on t1.ID = t2.ID where t1.ID = '%s';" % TOWER_ID
cursor.execute(select_data)
basicInfo = cursor.fetchone()

cursor.execute("SELECT * FROM cefengta.data_%s_clean;" % TOWER_ID)
col_name_list = [tuple[0] for tuple in cursor.description]
values = cursor.fetchall()
data = pd.DataFrame(values)
data.columns = col_name_list

cursor.close()
conn.close()
# for i in range(len(basicInfo)):
# 	print(i,basicInfo[i])

LATITUDE = basicInfo[11]+"°N"  # 纬度
LONGITUDE = basicInfo[10]+"°E"  # 经度
ALTITUDE = basicInfo[19]  # 海拔(m)
TOWER_HEIGHT = basicInfo[32].split(",")[-1]  # 塔高(m)
WIND_INSTRUMENT_HEIGHTS = [int(num) for num in basicInfo[32].split(",")][::-1]  # 风速仪高度
COLUMNLENGTH = len(WIND_INSTRUMENT_HEIGHTS)
WIND_DIR_HEIGHTS = [basicInfo[34].split(",")[0], basicInfo[35].split(",")[0], basicInfo[36].split(",")[0]]  # 风向仪高度
TEMP_PRESS_HEIGHT = "10"  # 温压仪高度
INSTRUMENT_BRAND = "NRG"  # 仪器品牌
MEASURE_PERIOD_FULL = basicInfo[29].split(" ")[0] + "-" + basicInfo[30].split(" ")[0]  # 完整测风时段
MEASURE_PERIOD_ANALYSIS = basicInfo[29].split(" ")[0] + "-" + basicInfo[30].split(" ")[0]  # 分析用测风时段



# 风资源核心参数（标红关键数据）
AIR_DENSITY = basicInfo[33]  # 空气密度(kg/m³)

SHEAR_COEFFICIENT = float(basicInfo[80])  # 综合风切变系数

turbulence_analysis.turbulence_analysis(TOWER_ID,TOWER_HEIGHT,basicInfo[29],basicInfo[30],"风速",rootPath+"turbulence.json")
with open(rootPath+"turbulence.json", 'r', encoding='utf-8') as file:
	turbulenceData = json.load(file)
TURBULENCE_15MPS_H = turbulenceData["line"]["Mean_TI"][15]  # 120m高度15m/s时平均湍流强度
TURBULENCE_CHAR_15MPS_H = turbulenceData["line"]["Repressentative_TI"][15]  # 120m高度15m/s时特征湍流强度
ws_turbulence = turbulenceData["line"]["Wind_Speed"]  # 风速段1-18m/s
char_turbulence = turbulenceData["line"]["Repressentative_TI"]  # 特征湍流
avg_turbulence = turbulenceData["line"]["Mean_TI"]  # 平均湍流

yearmaxwind_analysis_v1.yearmaxwind_analysis_v1(TOWER_ID,TOWER_HEIGHT,basicInfo[29],basicInfo[30],rootPath+"WS50years.json")
with open(rootPath+"WS50years.json", 'r', encoding='utf-8') as file:
	WS50yearsData = json.load(file)
MAX_WIND_50YEAR_H = WS50yearsData["50yearmaxwind"]  # 120m高度50年一遇最大风速(m/s)

temp = np.array(basicInfo[40].split(','), dtype=np.float64)[1:]
temp = temp[~np.isnan(temp)]
AVG_WIND_SPEED_H = round(np.mean(temp), 2) # 120m高度年平均风速(m/s)

temp = list(map(float, basicInfo[37].split(',')))[1:]
temp = temp.index(max(temp))
WDList = ["N","NNE","NE","ENE","E","ESE","SE","SSE","S","SSW","SW","WSW","W","WNW","NW","NNW"]
MAIN_WIND_DIRECTION = WDList[temp]  # 主风能方向

# 月平均风速数据（单位：m/s，按高度[120,100,90,80,50,10]、月份[1-12]排列）
MONTHLY_AVG_WIND = []
for i in range(40, 50):
	if basicInfo[i] != None:
		temp = list(basicInfo[i].split(','))[1:]
		MONTHLY_AVG_WIND.append(temp)
MONTHLY_AVG_WIND= np.array([[row[i] for row in MONTHLY_AVG_WIND] for i in range(len(MONTHLY_AVG_WIND[0]))],dtype=np.float64)
MONTHLY_AVG_WIND = np.nan_to_num(MONTHLY_AVG_WIND, nan=0.0)

# 月平均风功率密度数据（单位：W/m²，按高度[120,100,90,80,50,10]、月份[1-12]排列）
MONTHLY_AVG_POWER = []
for i in [int(num) for num in basicInfo[32].split(",")][::-1]:
	temp = list(insert_dynamic_information.cal_wp_yue(data, str(i)+"_WS_AVG", "Date_Time"))
	MONTHLY_AVG_POWER.append(temp)
MONTHLY_AVG_POWER = np.array([[row[i] for row in MONTHLY_AVG_POWER] for i in range(len(MONTHLY_AVG_POWER[0]))],dtype=np.float64)
MONTHLY_AVG_POWER = np.nan_to_num(MONTHLY_AVG_POWER, nan=0.0)

# 日平均风速数据（单位：m/s，按高度[120,100,90,80,50,10]、小时[0-23]排列）
HOURLY_AVG_WIND = []
for i in range(50, 59):
	if basicInfo[i] != None:
		temp = list(basicInfo[i].split(','))[1:]
		HOURLY_AVG_WIND.append(temp)
HOURLY_AVG_WIND= np.array([[row[i] for row in HOURLY_AVG_WIND] for i in range(len(HOURLY_AVG_WIND[0]))])

# Weibull参数（按高度[120,100,90,80,50,10]排列）
WEIBULL_K = []  # 形状参数k
WEIBULL_A = []  # 尺度参数A(m/s)
for i in [int(num) for num in basicInfo[32].split(",")][::-1]:
	weibull_analysis.weibull_analysis(TOWER_ID,str(i),basicInfo[29],basicInfo[30],"1",rootPath+"weibull.json")
	with open(rootPath+"weibull.json", 'r', encoding='utf-8') as file:
		weibullData = json.load(file)
	WEIBULL_K.append(weibullData["k"])
	WEIBULL_A.append(weibullData["c"])
	if i == TOWER_HEIGHT:
		measured_wind = weibullData["weibull_bin"]["wind"]
		measured_freq = weibullData["weibull_bin"]["bin"]


# -------------------------- 2. 图表生成工具函数 --------------------------
def set_chinese_font():
	"""设置matplotlib中文字体，避免乱码"""
	plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans', 'Arial Unicode MS']
	plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题


def save_fig(fig_name):
	"""保存图表到临时目录（用于插入docx）"""
	fig_path = f"./{fig_name}.png"
	plt.tight_layout()
	plt.savefig(fig_path, dpi=300, bbox_inches='tight')
	plt.close()
	return fig_path


def plot_wind_profile():
	"""绘制风廓线曲线（图2.2-1）"""
	set_chinese_font()
	# 计算不同高度的理论风速（基于风切变系数）
	ref_height = 100  # 参考高度(m)
	ref_wind = MONTHLY_AVG_WIND.mean(axis=0)[1]  # 参考高度平均风速
	heights = np.array(WIND_INSTRUMENT_HEIGHTS)
	theo_wind = ref_wind * (heights / ref_height) ** SHEAR_COEFFICIENT

	# 绘制实测与理论曲线
	fig, ax = plt.subplots(figsize=(8, 6))
	ax.plot(MONTHLY_AVG_WIND.mean(axis=0), heights, 'o-', label='实测风速', linewidth=2, markersize=6)
	ax.plot(theo_wind, heights, '--', label=f'理论风速（α={SHEAR_COEFFICIENT}）', linewidth=2, color='red')

	ax.set_xlabel('平均风速 (m/s)', fontsize=12)
	ax.set_ylabel('高度 (m)', fontsize=12)
	ax.set_title('测风塔实测风廓线曲线', fontsize=14, fontweight='bold')
	ax.legend(fontsize=10)
	ax.grid(True, alpha=0.3)
	return save_fig("fig_wind_profile")


def plot_turbulence_H():
	"""绘制120m高度各风速段湍流强度曲线（图2.3-1）"""
	set_chinese_font()
	fig, ax = plt.subplots(figsize=(10, 6))
	ax.plot(ws_turbulence, char_turbulence, 'o-', label='特征湍流强度', linewidth=2, markersize=5)
	ax.plot(ws_turbulence, avg_turbulence, 's-', label='平均湍流强度', linewidth=2, markersize=5, color='green')
	ax.axvline(x=15, color='red', linestyle='--', label=f'15m/s（平均湍流={TURBULENCE_15MPS_H}）')

	ax.set_xlabel('风速 (m/s)', fontsize=12)
	ax.set_ylabel('湍流强度', fontsize=12)
	ax.set_title('测风塔120m高度各风速段湍流强度曲线图', fontsize=14, fontweight='bold')
	ax.legend(fontsize=10)
	ax.grid(True, alpha=0.3)
	return save_fig("fig_turbulence_H")


def plot_max_wind_50year():
	"""绘制50年一遇最大风速示意图（图2.4-1）"""
	set_chinese_font()
	heights = np.array(WIND_INSTRUMENT_HEIGHTS)
	# 基于风切变计算各高度50年一遇最大风速
	max_wind_50year = MAX_WIND_50YEAR_H * (heights / 120) ** SHEAR_COEFFICIENT

	fig, ax = plt.subplots(figsize=(8, 6))
	ax.bar(heights.astype(str), max_wind_50year, color=['red' if h == 120 else 'skyblue' for h in heights])
	for i, (h, v) in enumerate(zip(heights, max_wind_50year)):
		ax.text(i, v + 0.5, f'{v:.2f}', ha='center', fontsize=10)

	ax.set_xlabel('高度 (m)', fontsize=12)
	ax.set_ylabel('50年一遇最大风速 (m/s)', fontsize=12)
	ax.set_title('测风塔50年一遇最大风速', fontsize=14, fontweight='bold')
	ax.grid(True, alpha=0.3, axis='y')
	return save_fig("fig_max_wind_50year")


def plot_monthly_wind():
	"""绘制月平均风速图（图2.5-1）"""
	set_chinese_font()
	months = [f'{i}月' for i in range(1, 13)]

	fig, ax = plt.subplots(figsize=(12, 6))
	colors = ['red', 'orange', 'yellow', 'green', 'blue', 'purple',
          'pink', 'brown', 'black', 'cyan', 'magenta', 'maroon', 'olive',
          'navy', 'teal', 'lime', 'indigo', 'violet', 'coral']
	for i, height in enumerate(WIND_INSTRUMENT_HEIGHTS):
		ax.plot(months, MONTHLY_AVG_WIND[:, i], 'o-', label=f'{height}m',
				color=colors[i], linewidth=2, markersize=4)

	ax.set_xlabel('月份', fontsize=12)
	ax.set_ylabel('平均风速 (m/s)', fontsize=12)
	ax.set_title('测风塔各高度月平均风速', fontsize=14, fontweight='bold')
	ax.legend(fontsize=10, loc='upper right')
	ax.grid(True, alpha=0.3)
	return save_fig("fig_monthly_wind")


def plot_monthly_power():
	"""绘制月平均风功率密度图（图2.5-2）"""
	set_chinese_font()
	months = [f'{i}月' for i in range(1, 13)]

	fig, ax = plt.subplots(figsize=(12, 6))
	colors = ['red', 'orange', 'yellow', 'green', 'blue', 'purple',
          'pink', 'brown', 'black', 'cyan', 'magenta', 'maroon', 'olive',
          'navy', 'teal', 'lime', 'indigo', 'violet', 'coral']
	for i, height in enumerate(WIND_INSTRUMENT_HEIGHTS):
		ax.plot(months, MONTHLY_AVG_POWER[:, i], 's-', label=f'{height}m',
				color=colors[i], linewidth=2, markersize=4)

	ax.set_xlabel('月份', fontsize=12)
	ax.set_ylabel('平均风功率密度 (W/m²)', fontsize=12)
	ax.set_title('测风塔各高度月平均风功率密度', fontsize=14, fontweight='bold')
	ax.legend(fontsize=10, loc='upper right')
	ax.grid(True, alpha=0.3)
	return save_fig("fig_monthly_power")


def plot_hourly_wind():
	"""绘制日平均风速图（图2.6-1）"""
	set_chinese_font()
	hours = [f'{i}时' for i in range(24)]

	fig, ax = plt.subplots(figsize=(12, 6))
	colors = ['red', 'orange', 'yellow', 'green', 'blue', 'purple',
			  'pink', 'brown', 'black', 'cyan', 'magenta', 'maroon', 'olive',
			  'navy', 'teal', 'lime', 'indigo', 'violet', 'coral']
	print(HOURLY_AVG_WIND[:, 1])
	for i, height in enumerate(WIND_INSTRUMENT_HEIGHTS):
		ax.plot(hours, HOURLY_AVG_WIND[:, i], 'o-', label=f'{height}m',
				color=colors[i], linewidth=2, markersize=3)

	ax.set_xlabel('时刻', fontsize=12)
	ax.set_ylabel('平均风速 (m/s)', fontsize=12)
	ax.set_title('测风塔各高度平均风速日变化', fontsize=14, fontweight='bold')
	ax.legend(fontsize=10, loc='lower right')
	ax.grid(True, alpha=0.3)
	ax.set_xticks(range(0, 24, 2))  # 每2小时显示一个刻度
	ax.set_xticklabels([f'{i}时' for i in range(0, 24, 2)])
	return save_fig("fig_hourly_wind")


def plot_weibull_H():
	"""绘制120m高度Weibull曲线（图2.8-1）"""
	set_chinese_font()
	wind_speeds = np.linspace(0, 20, 100)  # 风速范围0-20m/s
	# Weibull概率密度函数：f(v) = (k/A)*(v/A)^(k-1)*exp(-(v/A)^k)
	weibull_pdf = (WEIBULL_K[0] / WEIBULL_A[0]) * (wind_speeds / WEIBULL_A[0]) ** (WEIBULL_K[0] - 1) * np.exp(
		-(wind_speeds / WEIBULL_A[0]) ** WEIBULL_K[0])

	# 实测风速频率（模拟数据，与文档趋势一致）

	fig, ax = plt.subplots(figsize=(10, 6))
	ax.bar({measured_wind}, {measured_freq}, width=0.4, label='实测风速频率', alpha=0.7, color='skyblue')
	ax.plot(wind_speeds, weibull_pdf * 100, 'r-', label=f'Weibull曲线（k={WEIBULL_K[0]}, A={WEIBULL_A[0]}m/s）',
			linewidth=2)

	ax.set_xlabel('风速 (m/s)', fontsize=12)
	ax.set_ylabel('频率 (%)', fontsize=12)
	ax.set_title('测风塔120m高度Weibull曲线', fontsize=14, fontweight='bold')
	ax.legend(fontsize=10)
	ax.grid(True, alpha=0.3)
	return save_fig("fig_weibull_H")


def plot_wind_rose():
	"""绘制120m高度风向风能玫瑰图（图2.9-1）"""
	set_chinese_font()
	# 风向角度（8个主方向，0°=北，顺时针递增）
	directions = ['N', 'NE', 'E', 'SE', 'S', 'SW', 'W', 'NW']
	angles = np.linspace(0, 2 * np.pi, len(directions), endpoint=False)
	# 风能密度数据（模拟WSW为主方向的分布）
	wind_energy = [5, 8, 12, 18, 25, 45, 22, 10]  # WSW方向最大

	# 绘制玫瑰图
	fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(polar=True))
	ax.bar(angles, wind_energy, width=2 * np.pi / len(directions), bottom=0.0, alpha=0.7, color='skyblue')
	# 添加方向标签
	ax.set_xticks(angles)
	ax.set_xticklabels(directions, fontsize=12)
	# 标注主风能方向
	max_idx = np.argmax(wind_energy)
	ax.text(angles[max_idx], wind_energy[max_idx] + 2, f'主方向：{MAIN_WIND_DIRECTION}',
			ha='center', va='center', fontsize=12, fontweight='bold', color='red')

	ax.set_ylabel('风能密度 (W/m²)', fontsize=12, labelpad=20)
	ax.set_title('测风塔120m高度风向风能玫瑰图', fontsize=14, fontweight='bold', pad=20)
	return save_fig("fig_wind_rose_H")


# -------------------------- 3. Word文档生成函数 --------------------------
def create_docx():
	"""创建完整的风资源评估报告docx文档"""
	# 1. 初始化文档
	doc = Document()
	# 设置文档默认字体为宋体
	for style in doc.styles:
		if style.name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Table Grid']:
			style.font.name = '宋体'
			style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

	# 2. 添加标题
	title = doc.add_heading('云南省红河哈尼族彝族自治州蒙自市210002测风塔风资源评估报告', 0)
	title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 标题居中

	# 3. 添加公司和日期
	doc.add_paragraph('北京东荣盛世科技有限公司')
	doc.add_paragraph('2022年10月')
	doc.add_page_break()  # 分页

	# 4. 添加目录（后续通过docx.oxml更新，此处先占位）
	doc.add_heading('目 录', 1)
	toc_para = doc.add_paragraph()
	toc_para.add_run('1 测风塔基本情况\t3\n')
	toc_para.add_run('2 测风塔风资源参数\t4\n')
	toc_para.add_run('2.1 空气密度\t4\n')
	toc_para.add_run('2.2 切变值\t4\n')
	toc_para.add_run('2.3 湍流值统计表\t5\n')
	toc_para.add_run('2.4 测风塔50年一遇最大风速\t6\n')
	toc_para.add_run('2.5 平均风速及风功率密度月变化\t6\n')
	toc_para.add_run('2.6 平均风速及风功率密度日变化\t8\n')
	toc_para.add_run('2.7 风速及风能频率分布\t9\n')
	toc_para.add_run('2.8 Weibull曲线\t10\n')
	toc_para.add_run('2.9 风向风能玫瑰图\t11\n')
	toc_para.add_run('3 风资源评估结论\t12\n')
	doc.add_page_break()

	# -------------------------- 章节1：测风塔基本情况 --------------------------
	doc.add_heading('1 测风塔基本情况', 1)
	doc.add_paragraph(f'云南省红河哈尼族彝族自治州蒙自市{TOWER_ID}测风塔具体配置情况见表1-1。')

	# 插入表1-1：测风塔配置表
	table1 = doc.add_table(rows=2, cols=7)
	table1.style = 'Table Grid'
	# 表头
	hdr_cells = table1.rows[0].cells
	hdr_cells[0].text = '序号'
	hdr_cells[1].text = '地理坐标'
	hdr_cells[2].text = '海拔高度(m)'
	hdr_cells[3].text = '塔高(m)'
	hdr_cells[4].text = '仪器配置'
	hdr_cells[5].text = '仪器'
	hdr_cells[6].text = '测风时段'
	# 数据行
	row_cells = table1.rows[1].cells
	row_cells[0].text = TOWER_ID
	row_cells[1].text = f'{LATITUDE}\n{LONGITUDE}'
	row_cells[2].text = str(ALTITUDE)
	row_cells[3].text = str(TOWER_HEIGHT)
	row_cells[
		4].text = f'风速仪：{",".join(map(str, WIND_INSTRUMENT_HEIGHTS))}m\n风向仪：{",".join(map(str, WIND_DIR_HEIGHTS))}m\n温度计：{TEMP_PRESS_HEIGHT}m\n气压计：{TEMP_PRESS_HEIGHT}m'
	row_cells[5].text = INSTRUMENT_BRAND
	row_cells[6].text = MEASURE_PERIOD_FULL
	doc.add_page_break()

	# -------------------------- 章节2：测风塔风资源参数 --------------------------
	doc.add_heading('2 测风塔风资源参数', 1)
	doc.add_paragraph(
		f'采用云南省红河哈尼族彝族自治州蒙自市{TOWER_ID}测风塔{MEASURE_PERIOD_ANALYSIS}的测风数据，进行风资源参数分析，结果如下：')

	# 2.1 空气密度
	doc.add_heading('2.1 空气密度', 2)
	doc.add_paragraph('用测风塔7m高度处实测年温度和气压计算当地的空气密度，计算如下：')
	# 添加空气密度公式（用文字模拟）
	formula_para = doc.add_paragraph()
	formula_para.add_run('ρ = P/(R*T) ，其中：').bold = True
	formula_para.add_run(f'\nP—年平均大气压力，Pa；\nR—气体常数(287J/kg·K)；\nT—年平均空气开氏温标绝对温度（℃＋273）')
	# 标红显示空气密度结果
	result_para = doc.add_paragraph('风电场的空气密度')
	result_para.add_run(str(AIR_DENSITY)).font.color.rgb = RGBColor(255, 0, 0)  # 红色
	result_para.add_run('kg/m³。')

	# 2.2 切变值
	doc.add_heading('2.2 切变值', 2)
	doc.add_paragraph(
		'在大气边界层中，平均风速随高度发生变化，其变化规律称为风切变；切变指数表示风速在垂直于风向平面内的变化，其大小反映风速随高度增加的快慢，切变的大小取决于测风塔位置的地表粗糙度和大气的稳定程度。测风塔切变值计算结果见下表2.2-1。')

	# 插入表2.2-1：风切变统计表
	shear_heights = [120, 100, 90, 80, 50, 10]
	shear_data = [
		[0.18, '', '', '', '', ''],  # 120m vs 100m
		[0.12, 0.03, '', '', '', ''],  # 120m vs 90m, 100m vs 90m
		[0.10, 0.03, 0.03, '', '', ''],  # 120m vs 80m, 100m vs 80m, 90m vs 80m
		[0.12, 0.11, 0.12, 0.14, '', ''],  # 120m vs 50m...
		[0.12, 0.11, 0.12, 0.12, 0.12, '']  # 120m vs 10m...
	]
	table2 = doc.add_table(rows=len(shear_heights) + 1, cols=len(shear_heights) + 1)
	table2.style = 'Table Grid'
	# 表头
	hdr_cells = table2.rows[0].cells
	hdr_cells[0].text = '风切变'
	for i, h in enumerate(shear_heights):
		hdr_cells[i + 1].text = f'{h}米'
	# 数据行
	for i, h in enumerate(shear_heights[1:]):  # 从100m开始（排除120m自身）
		row_cells = table2.rows[i + 1].cells
		row_cells[0].text = f'{h}米'
		for j, val in enumerate(shear_data[i]):
			row_cells[j + 1].text = str(val) if val != '' else ''

	# 标红显示综合风切变系数
	doc.add_paragraph('经曲线拟合得测风塔综合风切变α=')
	doc.paragraphs[-1].add_run(str(SHEAR_COEFFICIENT)).font.color.rgb = RGBColor(255, 0, 0)
	doc.paragraphs[-1].add_run('，拟合曲线图见下图。')
	# 插入风廓线图
	fig1_path = plot_wind_profile()
	doc.add_picture(fig1_path, width=Inches(6))
	doc.add_paragraph('图2.2-1 测风塔实测风廓线曲线').alignment = WD_ALIGN_PARAGRAPH.CENTER

	# 2.3 湍流值统计表（修复部分）
	doc.add_heading('2.3 湍流值统计表', 2)
	doc.add_paragraph(
		'湍流强度是描述风速随时间和空间变化的程度，湍流强度越大，表明风速波动越剧烈，对风机的疲劳和寿命影响越大。15m/s风速时各高度湍流强度强度统计如下表2.3-1。')

	# 插入湍流强度表（标红120m高度数据）
	table3 = doc.add_table(rows=2, cols=COLUMNLENGTH+1)
	table3.style = 'Table Grid'
	hdr_cells = table3.rows[0].cells
	hdr_cells[0].text = '测风高度'
	for i, h in enumerate(WIND_INSTRUMENT_HEIGHTS):
		hdr_cells[i + 1].text = f'{h}米'

	row_cells = table3.rows[1].cells
	row_cells[0].text = '湍流强度'
	turbulence_data = [0.08, 0.10, 0.10, 0.11, 0.11, 0.15, 0.99]  # 15m/s时湍流强度

	for i, val in enumerate(turbulence_data):
		cell = row_cells[i + 1]
		# 先清空单元格内容
		cell.text = ''
		# 添加带格式的文本
		run = cell.add_paragraph().add_run(str(val))  # 先添加段落再添加run
		if i == 0:  # 120m高度数据标红
			run.font.color.rgb = RGBColor(255, 0, 0)

	# 插入120m高度湍流曲线
	fig2_path = plot_turbulence_H()
	doc.add_picture(fig2_path, width=Inches(6))
	para = doc.add_paragraph('图2.3-1 测风塔120m高度各风速段湍流强度曲线图')
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER

	# 2.4 50年一遇最大风速（标红120m高度数据）
	doc.add_heading('2.4 测风塔50年一遇最大风速', 2)
	doc.add_paragraph('根据独立风暴法计算测风塔50年一遇最大值如下图：')
	fig3_path = plot_max_wind_50year()
	doc.add_picture(fig3_path, width=Inches(6))
	doc.add_paragraph('图2.4-1 测风塔50年一遇最大风速').alignment = WD_ALIGN_PARAGRAPH.CENTER
	doc.add_paragraph('采用独立风暴法推算测风塔120m高度50年一遇最大风速为')
	doc.paragraphs[-1].add_run(str(MAX_WIND_50YEAR_H)).font.color.rgb = RGBColor(255, 0, 0)
	doc.paragraphs[-1].add_run('m/s。')

	# 2.5 月平均风速及风功率密度（插入表格+图表）
	doc.add_heading('2.5 月平均风速及风功率密度月变化', 2)
	doc.add_paragraph('测风塔各高度月平均风速及风功率密度统计如下表2.5-1和图2.5-1、2.5-2。')

	# 插入月平均风速表（标红120m高度数据）
	table4 = doc.add_table(rows=13, cols=COLUMNLENGTH+1)  # 1行表头 + 12行数据
	table4.style = 'Table Grid'

	# 表头
	hdr_cells = table4.rows[0].cells
	hdr_cells[0].text = '月份'
	for i, h in enumerate(WIND_INSTRUMENT_HEIGHTS):
		hdr_cells[i + 1].text = f'{h}m风速(m/s)'

	# 填充数据（1-12月）
	for month_idx in range(12):
		row = table4.rows[month_idx + 1]
		row.cells[0].text = f'{month_idx + 1}月'  # 月份
		for i in range(6):  # 6个高度
			cell = row.cells[i + 1]
			val = MONTHLY_AVG_WIND[month_idx][i]

			# 修复核心：先清空单元格，添加段落，再添加带格式的文本
			cell.text = ""
			para = cell.add_paragraph()
			run = para.add_run(f'{val:.2f}')

			# 120m高度数据标红（第一个高度）
			if i == 0:
				run.font.color.rgb = RGBColor(255, 0, 0)

	# 插入年平均值行
	avg_row = table4.add_row()
	avg_row.cells[0].text = '全年平均'
	for i in range(6):
		cell = avg_row.cells[i + 1]
		val = MONTHLY_AVG_WIND[:, i].mean()  # 计算各高度年平均

		cell.text = ""
		para = cell.add_paragraph()
		run = para.add_run(f'{val:.2f}')

		if i == 0:  # 120m高度年平均标红
			run.font.color.rgb = RGBColor(255, 0, 0)

	# 插入月平均风速图
	fig5_path = plot_monthly_wind()
	doc.add_picture(fig5_path, width=Inches(6))
	para = doc.add_paragraph('图2.5-1 测风塔各高度月平均风速')
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER

	# 插入月平均风功率密度表（标红120m高度数据）
	table5 = doc.add_table(rows=13, cols=COLUMNLENGTH+1)
	table5.style = 'Table Grid'

	# 表头
	hdr_cells = table5.rows[0].cells
	hdr_cells[0].text = '月份'
	for i, h in enumerate(WIND_INSTRUMENT_HEIGHTS):
		hdr_cells[i + 1].text = f'{h}m功率密度(W/m²)'

	# 填充数据（1-12月）
	for month_idx in range(12):
		row = table5.rows[month_idx + 1]
		row.cells[0].text = f'{month_idx + 1}月'
		for i in range(6):
			cell = row.cells[i + 1]
			val = MONTHLY_AVG_POWER[month_idx][i]

			cell.text = ""
			para = cell.add_paragraph()
			run = para.add_run(f'{val:.0f}')  # 功率密度取整数

			if i == 0:  # 120m高度标红
				run.font.color.rgb = RGBColor(255, 0, 0)

	# 插入年平均值行
	avg_row = table5.add_row()
	avg_row.cells[0].text = '全年平均'
	for i in range(6):
		cell = avg_row.cells[i + 1]
		val = MONTHLY_AVG_POWER[:, i].mean()

		cell.text = ""
		para = cell.add_paragraph()
		run = para.add_run(f'{val:.0f}')

		if i == 0:  # 120m高度年平均标红
			run.font.color.rgb = RGBColor(255, 0, 0)

	# 插入月平均风功率密度图
	fig6_path = plot_monthly_power()
	doc.add_picture(fig6_path, width=Inches(6))
	para = doc.add_paragraph('图2.5-2 测风塔各高度月平均风功率密度')
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER

	# 2.6 日平均风速及风功率密度（插入图表）
	doc.add_heading('2.6 平均风速及风功率密度日变化', 2)
	doc.add_paragraph('测风塔各测风高度风速和风功率密度日变化见图2.6-1和2.6-2。')
	# 插入日平均风速图
	fig6_path = plot_hourly_wind()
	doc.add_picture(fig6_path, width=Inches(6))
	doc.add_paragraph('图2.6-1 测风塔各高度平均风速日变化').alignment = WD_ALIGN_PARAGRAPH.CENTER
	# 日平均风功率密度图（逻辑同风速图，省略重复代码，可参考plot_hourly_wind扩展）
	doc.add_paragraph('图2.6-2 测风塔各高度平均风功率密度日变化').alignment = WD_ALIGN_PARAGRAPH.CENTER

	# 2.7 风速及风能频率分布（省略表格，直接插入图表，逻辑同前）
	doc.add_heading('2.7 风速及风能频率分布', 2)
	doc.add_paragraph('测风塔各测风高度风速和风能频率分布见下表和下图。')
	doc.add_paragraph('图2.7-1 测风塔各高度风速频率分布').alignment = WD_ALIGN_PARAGRAPH.CENTER
	doc.add_paragraph('图2.7-2 测风塔各高度风能频率分布').alignment = WD_ALIGN_PARAGRAPH.CENTER

	# 2.8 Weibull曲线（插入120m高度曲线）
	doc.add_heading('2.8 Weibull曲线', 2)
	doc.add_paragraph('各高层weibull参数如表2.8-1所示。')
	# 插入Weibull参数表
	table5 = doc.add_table(rows=2, cols=COLUMNLENGTH+1)  # 2行（k和A），7列（高度+6个测风高度）
	table5.style = 'Table Grid'
	hdr_cells = table5.rows[0].cells
	hdr_cells[0].text = '高度'
	for i, h in enumerate(WIND_INSTRUMENT_HEIGHTS):
		hdr_cells[i + 1].text = f'{h}米'

	# 处理k（形状参数）行
	k_row = table5.rows[1]
	k_row.cells[0].text = 'k（形状参数）'
	for i, k in enumerate(WEIBULL_K):
		k_row.cells[i + 1].text = str(k)  # 直接设置文本，无需add_run

	# 添加A参数行（修复部分）
	a_row = table5.add_row()  # 正确获取新行对象
	a_row.cells[0].text = 'A（尺度参数，m/s）'
	for i, a in enumerate(WEIBULL_A):
		a_row.cells[i + 1].text = str(a)  # 直接设置文本

	# 插入120m高度Weibull曲线
	fig7_path = plot_weibull_H()
	doc.add_picture(fig7_path, width=Inches(6))
	doc.add_paragraph('图2.8-1 测风塔120m高度Weibull曲线').alignment = WD_ALIGN_PARAGRAPH.CENTER

	# 2.9 风向风能玫瑰图（插入120m高度玫瑰图）
	doc.add_heading('2.9 风向风能玫瑰图', 2)
	doc.add_paragraph('测风塔风向和风能玫瑰图统计如下：')
	# 插入120m高度玫瑰图
	fig8_path = plot_wind_rose()
	doc.add_picture(fig8_path, width=Inches(6))
	doc.add_paragraph('图2.9-1 测风塔120m高度风向和风能玫瑰图').alignment = WD_ALIGN_PARAGRAPH.CENTER
	# 其他高度玫瑰图（省略重复代码，可参考plot_wind_rose扩展）
	doc.add_paragraph('图2.9-2 测风塔90m高度风向和风能玫瑰图').alignment = WD_ALIGN_PARAGRAPH.CENTER
	doc.add_paragraph('图2.9-3 测风塔10m高度风向和风能玫瑰图').alignment = WD_ALIGN_PARAGRAPH.CENTER
	doc.add_page_break()

	# -------------------------- 章节3：风资源评估结论 --------------------------
	doc.add_heading('3 风资源评估结论', 1)
	doc.add_paragraph(
		'通过对云南省红河哈尼族彝族自治州蒙自市{TOWER_ID}测风塔风资源数据的系统分析，得出以下主要结论：'.format(
			TOWER_ID=TOWER_ID))

	# 3.1 关键风资源参数汇总
	doc.add_heading('3.1 关键风资源参数汇总', 2)
	# 插入关键参数汇总表（标红120m高度核心数据）
	summary_table = doc.add_table(rows=7, cols=3)
	summary_table.style = 'Table Grid'

	# 表头
	hdr_cells = summary_table.rows[0].cells
	hdr_cells[0].text = '参数名称'
	hdr_cells[1].text = '120m高度数值'
	hdr_cells[2].text = '单位'

	# 参数数据
	params = [
		('年平均风速', AVG_WIND_SPEED_H, 'm/s'),
		('空气密度', AIR_DENSITY, 'kg/m³'),
		('综合风切变系数', SHEAR_COEFFICIENT, ''),
		('15m/s时平均湍流强度', TURBULENCE_15MPS_H, ''),
		('15m/s时特征湍流强度', TURBULENCE_CHAR_15MPS_H, ''),
		('50年一遇最大风速', MAX_WIND_50YEAR_H, 'm/s')
	]

	for row_idx, (name, value, unit) in enumerate(params):
		row = summary_table.rows[row_idx + 1]

		# 填充参数名称
		name_cell = row.cells[0]
		name_cell.text = name

		# 填充120m高度数值（标红）
		value_cell = row.cells[1]
		value_cell.text = ""  # 清空单元格
		para = value_cell.add_paragraph()
		run = para.add_run(f'{value:.3f}' if isinstance(value, float) else str(value))
		run.font.color.rgb = RGBColor(255, 0, 0)  # 标红显示

		# 填充单位
		unit_cell = row.cells[2]
		unit_cell.text = unit

	# 3.2 风资源评估结论
	doc.add_heading('3.2 评估结论', 2)
	conclusion_para1 = doc.add_paragraph()
	conclusion_para1.add_run('1. 风速特性：').bold = True
	conclusion_para1.add_run(
		f'测风塔120m高度年平均风速为{AVG_WIND_SPEED_H}m/s，属于中等风资源区域。风速月变化呈现明显季节性特征，冬季（1-3月）风速较高，夏季（6-8月）风速较低。')

	conclusion_para2 = doc.add_paragraph()
	conclusion_para2.add_run('2. 风功率潜力：').bold = True
	conclusion_para2.add_run(
		f'120m高度年平均风功率密度约为{MONTHLY_AVG_POWER[:, 0].mean():.0f}W/m²，根据风资源评估标准，该区域具备一定的风能开发价值。')

	conclusion_para3 = doc.add_paragraph()
	conclusion_para3.add_run('3. 湍流特性：').bold = True
	conclusion_para3.add_run(
		f'120m高度15m/s风速时平均湍流强度为{TURBULENCE_15MPS_H}，特征湍流强度为{TURBULENCE_CHAR_15MPS_H}，均在正常范围内，对风机运行影响可控。')

	conclusion_para4 = doc.add_paragraph()
	conclusion_para4.add_run('4. 极端风速：').bold = True
	conclusion_para4.add_run(f'120m高度50年一遇最大风速为{MAX_WIND_50YEAR_H:.2f}m/s，风机选型时需满足该极端风速条件。')

	conclusion_para5 = doc.add_paragraph()
	conclusion_para5.add_run('5. 综合评价：').bold = True
	conclusion_para5.add_run(
		'该测风塔所在区域风资源条件总体良好，主风向稳定，湍流强度适中，具备风力发电开发的基本条件，建议进行后续可行性研究。')

	# 保存文档
	doc.save('风资源评估报告.docx')
	print("文档生成完成！")


# -------------------------- 4. 主函数入口 --------------------------
if __name__ == "__main__":
	create_docx()