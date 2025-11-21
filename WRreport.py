from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.enum.section import WD_SECTION
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import warnings
import pymysql
import json
from datetime import datetime
from function import turbulence_analysis, yearmaxwind_analysis_v1, insert_dynamic_information, weibull_analysis, \
	daily_analysis, frequency_analysis, shear_analysis
from report import turbulenceData

warnings.filterwarnings('ignore')
host = 'localhost'
port = 3306
user = 'root'  # 用户名
password = '123456'  # 密码
database = 'cefengta'  # 数据库名称
rootPath = "C:\\Users\\admin\\PycharmProjects\\metaReport\\"

# -------------------------- 1. 核心参数定义 --------------------------
# 测风塔基础信息
TOWER_ID = "14648"

conn = pymysql.connect(host=host, port=port, user=user, password=password, database=database, charset='utf8mb4')
cursor = conn.cursor()
select_data = "SELECT * FROM cefengta.static_information as t1 inner join cefengta.dynamic_information as t2 on t1.ID = t2.ID where t1.ID = '%s';" % TOWER_ID
cursor.execute(select_data)
basicInfo = cursor.fetchone()

cursor.execute("SELECT * FROM cefengta.data_%s_clean;" % TOWER_ID)
col_name_list = [tuple[0] for tuple in cursor.description]
values = cursor.fetchall()
data = pd.DataFrame(values)
data.columns = col_name_list

cursor.close()
conn.close()

TOWER_NAME = basicInfo[1]
LOCATION = basicInfo[7] + basicInfo[8] + basicInfo[9]
LATITUDE = basicInfo[11] + "°N"  # 纬度
LONGITUDE = basicInfo[10] + "°E"  # 经度
ALTITUDE = basicInfo[19]  # 海拔(m)
TOWER_HEIGHT = basicInfo[32].split(",")[-1]  # 塔高(m)
WIND_INSTRUMENT_HEIGHTS = [int(num) for num in basicInfo[32].split(",")][::-1]  # 风速仪高度
COLUMNLENGTH = len(WIND_INSTRUMENT_HEIGHTS)
WIND_DIR_HEIGHTS = [basicInfo[34].split(",")[0], basicInfo[35].split(",")[0], basicInfo[36].split(",")[0]]  # 风向仪高度
TEMP_PRESS_HEIGHT = "10"  # 温压仪高度
INSTRUMENT_BRAND = "NRG"  # 仪器品牌
MEASURE_PERIOD_FULL = basicInfo[29].split(" ")[0] + "~" + basicInfo[30].split(" ")[0]  # 完整测风时段
MEASURE_PERIOD_ANALYSIS = basicInfo[29].split(" ")[0] + "~" + basicInfo[30].split(" ")[0]  # 分析用测风时段

# 风资源核心参数
AIR_DENSITY = basicInfo[33]  # 空气密度(kg/m³)

SHEAR_COEFFICIENT = float(basicInfo[80])  # 综合风切变系数

turbulence_data = []
for i in [int(num) for num in basicInfo[32].split(",")][::-1]:
	turbulence_analysis.turbulence_analysis(TOWER_ID, str(i), basicInfo[29], basicInfo[30], "风速",
											rootPath + "turbulence.json")
	with open(rootPath + "turbulence.json", 'r', encoding='utf-8') as file:
		turbulenceData = json.load(file)
	turbulence_data.append((turbulenceData["line"]["Mean_TI"][15]))
	if str(i) == str(TOWER_HEIGHT):
		TURBULENCE_15MPS_H = turbulenceData["line"]["Mean_TI"][15]  # TOWER_HEIGHTm高度15m/s时平均湍流强度
		TURBULENCE_CHAR_15MPS_H = turbulenceData["line"]["Repressentative_TI"][15]  # TOWER_HEIGHTm高度15m/s时代表湍流强度
		ws_turbulence = turbulenceData["line"]["Wind_Speed"]  # 风速段1-18m/s
		char_turbulence = turbulenceData["line"]["Repressentative_TI"]  # 代表湍流
		avg_turbulence = turbulenceData["line"]["Mean_TI"]  # 平均湍流

yearmaxwind_analysis_v1.yearmaxwind_analysis_v1(TOWER_ID, TOWER_HEIGHT, basicInfo[29], basicInfo[30],
												rootPath + "WS50years.json")
with open(rootPath + "WS50years.json", 'r', encoding='utf-8') as file:
	WS50yearsData = json.load(file)
MAX_WIND_50YEAR_H = WS50yearsData["50yearmaxwind"]  # TOWER_HEIGHTm高度50年一遇最大风速(m/s)

temp = np.array(basicInfo[40].split(','), dtype=np.float64)[1:]
temp = temp[~np.isnan(temp)]
AVG_WIND_SPEED_H = round(np.mean(temp), 2)  # TOWER_HEIGHTm高度年平均风速(m/s)

temp = list(map(float, basicInfo[37].split(',')))[1:]
temp = temp.index(max(temp))
WDList = ["N", "NNE", "NE", "ENE", "E", "ESE", "SE", "SSE", "S", "SSW", "SW", "WSW", "W", "WNW", "NW", "NNW"]
MAIN_WIND_DIRECTION = WDList[temp]  # 主风能方向

# 月平均风速数据
MONTHLY_AVG_WIND = []
for i in range(40, 50):
	if basicInfo[i] != None:
		temp = list(basicInfo[i].split(','))[1:]
		MONTHLY_AVG_WIND.append(temp)
MONTHLY_AVG_WIND = np.array([[row[i] for row in MONTHLY_AVG_WIND] for i in range(len(MONTHLY_AVG_WIND[0]))],
							dtype=np.float64)
MONTHLY_AVG_WIND = np.nan_to_num(MONTHLY_AVG_WIND, nan=0.0)

# 月平均风功率密度数据
MONTHLY_AVG_POWER = []
for i in [int(num) for num in basicInfo[32].split(",")][::-1]:
	temp = list(insert_dynamic_information.cal_wp_yue(data, str(i) + "_WS_AVG", "Date_Time"))
	MONTHLY_AVG_POWER.append(temp)
MONTHLY_AVG_POWER = np.array([[row[i] for row in MONTHLY_AVG_POWER] for i in range(len(MONTHLY_AVG_POWER[0]))],
							 dtype=np.float64)
MONTHLY_AVG_POWER = np.nan_to_num(MONTHLY_AVG_POWER, nan=0.0)

# 日平均风速数据
HOURLY_AVG_WIND = []
for i in range(50, 60):
	if basicInfo[i] != None:
		temp = list(basicInfo[i].split(','))[1:]
		HOURLY_AVG_WIND.append(temp)
HOURLY_AVG_WIND = np.array([[row[i] for row in HOURLY_AVG_WIND] for i in range(len(HOURLY_AVG_WIND[0]))],
						   dtype=np.float64)

# 日均风功率数据
HOURLY_AVG_POWER = []
hours = [f"{h:02d}" for h in range(24)]
for i in [int(num) for num in basicInfo[32].split(",")][::-1]:
	daily_analysis.daily_analysis(TOWER_ID, str(i), basicInfo[29], basicInfo[30], "年况", "风功率",
								  rootPath + "daily.json")
	with open(rootPath + "daily.json", 'r', encoding='utf-8') as file:
		dailyData = json.load(file)
	hourly_values = []
	for hour in hours:
		if hour in dailyData["year"]:
			hourly_values.append(dailyData["year"][hour])
		else:
			hourly_values.append(0)
	HOURLY_AVG_POWER.append(hourly_values)
HOURLY_AVG_POWER = np.array([[row[i] for row in HOURLY_AVG_POWER] for i in range(len(HOURLY_AVG_POWER[0]))],
							dtype=np.float64)

# 频率数据
WIND_FREQ = []
ENERGY_FREQ = []
for i in [int(num) for num in basicInfo[32].split(",")][::-1]:
	frequency_analysis.frequency_analysis(TOWER_ID, str(i), basicInfo[29], basicInfo[30], rootPath + "frequency.json")
	with open(rootPath + "frequency.json", 'r', encoding='utf-8') as file:
		frequencyData = json.load(file)
	wind_items = [(float(speed), freq) for speed, freq in frequencyData["wind"].items()]
	wind_items.sort(key=lambda x: x[0])
	WIND_FREQ.append([freq for _, freq in wind_items])

	wpd_items = [(float(speed), wpd) for speed, wpd in frequencyData["WPD"].items()]
	wpd_items.sort(key=lambda x: x[0])
	ENERGY_FREQ.append([wpd for _, wpd in wpd_items])

WIND_FREQ = np.array(WIND_FREQ, dtype=np.float64)
ENERGY_FREQ = np.array(ENERGY_FREQ, dtype=np.float64)

# Weibull参数
WEIBULL_K = []  # 形状参数k
WEIBULL_A = []  # 尺度参数A(m/s)
for i in [int(num) for num in basicInfo[32].split(",")][::-1]:
	weibull_analysis.weibull_analysis(TOWER_ID, str(i), basicInfo[29], basicInfo[30], "1", rootPath + "weibull.json")
	with open(rootPath + "weibull.json", 'r', encoding='utf-8') as file:
		weibullData = json.load(file)
	WEIBULL_K.append(weibullData["k"])
	WEIBULL_A.append(weibullData["c"])
	if str(i) == str(TOWER_HEIGHT):
		measured_wind = weibullData["weibull_bin"]["wind"]
		measured_freq = weibullData["weibull_bin"]["bin"]

# 风切变
shear_analysis.shear_analysis(TOWER_ID, basicInfo[29], basicInfo[30], rootPath + "shear.json")
with open(rootPath + "shear.json", 'r', encoding='utf-8') as file:
	shearData = json.load(file)
shear_heights = [int(num) for num in basicInfo[32].split(",")][::-1]
shear_data = []
shear_row_data = []
for i in range(COLUMNLENGTH - 1):
	shear_row_data = []
	for j in range(COLUMNLENGTH - 1):
		shear_row_data.append(shearData["shear_ceng"][i][str(shear_heights[j])])
	shear_data.append(list(shear_row_data))

# 风玫瑰
WIND_FREQ_STRINGS = []
wd_heights = 0
for i in range(34, 37):
	if basicInfo[i] != None:
		WIND_FREQ_STRINGS.append(basicInfo[i])
		wd_heights = wd_heights + 1
ENERGY_FREQ_STRINGS = []
for i in range(34, 37):
	if basicInfo[i] != None:
		ENERGY_FREQ_STRINGS.append(basicInfo[i])


# -------------------------- 2. 图表生成工具函数 --------------------------
def set_chinese_font():
	"""设置matplotlib中文字体，避免乱码"""
	plt.rcParams['font.sans-serif'] = ['SimHei', 'DejaVu Sans', 'Arial Unicode MS']
	plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题


def save_fig(fig_name):
	"""保存图表到临时目录（用于插入docx）"""
	fig_path = f"./{fig_name}.png"
	plt.tight_layout()
	plt.savefig(fig_path, dpi=300, bbox_inches='tight')
	plt.close()
	return fig_path


def plot_wind_profile():
	"""绘制风廓线曲线（图2.2-1）"""
	set_chinese_font()
	ref_height = 100  # 参考高度(m)
	ref_wind = MONTHLY_AVG_WIND.mean(axis=0)[1]  # 参考高度平均风速
	heights = np.array(WIND_INSTRUMENT_HEIGHTS)
	theo_wind = ref_wind * (heights / ref_height) ** SHEAR_COEFFICIENT

	fig, ax = plt.subplots(figsize=(8, 6))
	ax.plot(MONTHLY_AVG_WIND.mean(axis=0), heights, 'o-', label='实测风速', linewidth=2, markersize=6)
	ax.plot(theo_wind, heights, '--', label=f'理论风速（α={SHEAR_COEFFICIENT}）', linewidth=2, color='red')

	ax.set_xlabel('平均风速 (m/s)', fontsize=12)
	ax.set_ylabel('高度 (m)', fontsize=12)
	ax.set_title('测风塔实测风廓线曲线', fontsize=14, fontweight='bold')
	ax.legend(fontsize=10)
	ax.grid(True, alpha=0.3)
	return save_fig("fig_wind_profile")


def plot_turbulence_H():
	"""绘制TOWER_HEIGHTm高度各风速段湍流强度曲线（图2.3-1）"""
	set_chinese_font()
	fig, ax = plt.subplots(figsize=(10, 6))
	ax.plot(ws_turbulence, char_turbulence, 'o-', label='代表湍流强度', linewidth=2, markersize=5)
	ax.plot(ws_turbulence, avg_turbulence, 's-', label='平均湍流强度', linewidth=2, markersize=5, color='green')
	ax.axvline(x=15, color='red', linestyle='--', label=f'15m/s（平均湍流={TURBULENCE_15MPS_H}）')

	ax.set_xlabel('风速 (m/s)', fontsize=12)
	ax.set_ylabel('湍流强度', fontsize=12)
	ax.set_title(f'测风塔{TOWER_HEIGHT}m高度各风速段湍流强度曲线图', fontsize=14, fontweight='bold')
	ax.legend(fontsize=10)
	ax.grid(True, alpha=0.3)
	return save_fig("fig_turbulence_H")


def plot_max_wind_50year():
	"""绘制50年一遇最大风速示意图（图2.4-1）"""
	set_chinese_font()
	heights = np.array(WIND_INSTRUMENT_HEIGHTS)
	max_wind_50year = MAX_WIND_50YEAR_H * (heights / float(TOWER_HEIGHT)) ** SHEAR_COEFFICIENT

	fig, ax = plt.subplots(figsize=(8, 6))
	ax.bar(heights.astype(str), max_wind_50year, color=['red' if h == TOWER_HEIGHT else 'skyblue' for h in heights])
	for i, (h, v) in enumerate(zip(heights, max_wind_50year)):
		ax.text(i, v + 0.5, f'{v:.2f}', ha='center', fontsize=10)

	ax.set_xlabel('高度 (m)', fontsize=12)
	ax.set_ylabel('50年一遇最大风速 (m/s)', fontsize=12)
	ax.set_title('测风塔50年一遇最大风速', fontsize=14, fontweight='bold')
	ax.grid(True, alpha=0.3, axis='y')
	return save_fig("fig_max_wind_50year")


def plot_monthly_wind():
	"""绘制月平均风速图（图2.5-1）"""
	set_chinese_font()
	months = [f'{i}月' for i in range(1, 13)]

	fig, ax = plt.subplots(figsize=(12, 6))
	colors = ['red', 'orange', 'yellow', 'green', 'blue', 'purple',
			  'pink', 'brown', 'black', 'cyan', 'magenta', 'maroon', 'olive',
			  'navy', 'teal', 'lime', 'indigo', 'violet', 'coral']
	for i, height in enumerate(WIND_INSTRUMENT_HEIGHTS):
		ax.plot(months, MONTHLY_AVG_WIND[:, i], 'o-', label=f'{height}m',
				color=colors[i], linewidth=2, markersize=4)

	ax.set_xlabel('月份', fontsize=12)
	ax.set_ylabel('平均风速 (m/s)', fontsize=12)
	ax.set_title('测风塔各高度月平均风速', fontsize=14, fontweight='bold')
	ax.legend(fontsize=10, loc='upper right')
	ax.grid(True, alpha=0.3)
	return save_fig("fig_monthly_wind")


def plot_monthly_power():
	"""绘制月平均风功率密度图（图2.5-2）"""
	set_chinese_font()
	months = [f'{i}月' for i in range(1, 13)]

	fig, ax = plt.subplots(figsize=(12, 6))
	colors = ['red', 'orange', 'yellow', 'green', 'blue', 'purple',
			  'pink', 'brown', 'black', 'cyan', 'magenta', 'maroon', 'olive',
			  'navy', 'teal', 'lime', 'indigo', 'violet', 'coral']
	for i, height in enumerate(WIND_INSTRUMENT_HEIGHTS):
		ax.plot(months, MONTHLY_AVG_POWER[:, i], 's-', label=f'{height}m',
				color=colors[i], linewidth=2, markersize=4)

	ax.set_xlabel('月份', fontsize=12)
	ax.set_ylabel('平均风功率密度 (W/m²)', fontsize=12)
	ax.set_title('测风塔各高度月平均风功率密度', fontsize=14, fontweight='bold')
	ax.legend(fontsize=10, loc='upper right')
	ax.grid(True, alpha=0.3)
	return save_fig("fig_monthly_power")


def plot_hourly_wind():
	"""绘制日平均风速图（图2.6-1）"""
	set_chinese_font()
	hours = [f'{i}时' for i in range(24)]

	fig, ax = plt.subplots(figsize=(12, 6))
	colors = ['red', 'orange', 'yellow', 'green', 'blue', 'purple',
			  'pink', 'brown', 'black', 'cyan', 'magenta', 'maroon', 'olive',
			  'navy', 'teal', 'lime', 'indigo', 'violet', 'coral']
	for i, height in enumerate(WIND_INSTRUMENT_HEIGHTS):
		ax.plot(hours, HOURLY_AVG_WIND[:, i], 'o-', label=f'{height}m',
				color=colors[i], linewidth=2, markersize=3)

	ax.set_xlabel('时刻', fontsize=12)
	ax.set_ylabel('平均风速 (m/s)', fontsize=12)
	ax.set_title('测风塔各高度平均风速日变化', fontsize=14, fontweight='bold')
	ax.legend(fontsize=10, loc='lower right')
	ax.grid(True, alpha=0.3)
	ax.set_xticks(range(0, 24, 2))
	ax.set_xticklabels([f'{i}时' for i in range(0, 24, 2)])
	return save_fig("fig_hourly_wind")


def plot_hourly_power():
	"""绘制日平均风功率密度图（图2.6-2）"""
	set_chinese_font()
	hours = range(24)
	fig, ax = plt.subplots(figsize=(12, 6))

	colors = ['red', 'orange', 'yellow', 'green', 'blue', 'purple',
			  'pink', 'brown', 'black', 'cyan', 'magenta', 'maroon', 'olive',
			  'navy', 'teal', 'lime', 'indigo', 'violet', 'coral']
	valid_count = min(len(WIND_INSTRUMENT_HEIGHTS), HOURLY_AVG_POWER.shape[1])

	for i in range(valid_count):
		height = WIND_INSTRUMENT_HEIGHTS[i]
		ax.plot(hours, HOURLY_AVG_POWER[:, i], 's-', label=f'{height}m',
				color=colors[i], linewidth=2, markersize=4)

	ax.set_xlabel('时刻', fontsize=12)
	ax.set_ylabel('平均风功率密度 (W/m²)', fontsize=12)
	ax.set_title('测风塔各高度平均风功率密度日变化', fontsize=14, fontweight='bold')

	ax.set_xticks(range(0, 24, 2))
	ax.set_xticklabels([f'{i}时' for i in range(0, 24, 2)])
	ax.set_ylim(bottom=0)
	if HOURLY_AVG_POWER.size > 0:
		ax.set_ylim(top=HOURLY_AVG_POWER.max() * 1.1)

	ax.legend(fontsize=10, loc='upper right')
	ax.grid(True, alpha=0.3, linestyle='--')

	return save_fig("fig_hourly_power")


def plot_wind_frequency():
	"""绘制各高度风速频率分布图（图2.7-1）"""
	set_chinese_font()
	wind_speeds = np.arange(0.5, 25.5, 0.5)
	fig, ax = plt.subplots(figsize=(14, 7))

	colors = ['red', 'orange', 'yellow', 'green', 'blue', 'purple',
			  'pink', 'brown', 'black', 'cyan', 'magenta', 'maroon', 'olive',
			  'navy', 'teal', 'lime', 'indigo', 'violet', 'coral']
	linestyles = ['-', '--', '-.', ':', '-', '--', '-.', ':', '-', '--']
	valid_count = min(len(WIND_INSTRUMENT_HEIGHTS), len(WIND_FREQ))

	for i in range(valid_count):
		height = WIND_INSTRUMENT_HEIGHTS[i]
		ax.plot(wind_speeds, WIND_FREQ[i], label=f'{height}m',
				color=colors[i], linestyle=linestyles[i], linewidth=2, markersize=1)

	ax.set_xlabel('风速 (m/s)', fontsize=12)
	ax.set_ylabel('频率 (%)', fontsize=12)
	ax.set_title('测风塔各高度风速频率分布', fontsize=14, fontweight='bold')

	ax.set_xticks(range(0, 26, 2))
	ax.set_xlim(0.5, 25)
	if WIND_FREQ.size > 0:
		ax.set_ylim(0, WIND_FREQ.max() * 1.1)

	ax.legend(fontsize=10, loc='upper right')
	ax.grid(True, alpha=0.3, linestyle='--')

	return save_fig("fig_wind_frequency")


def plot_energy_frequency():
	"""绘制各高度风能频率分布图（图2.7-2）"""
	set_chinese_font()
	wind_speeds = np.arange(0.5, 25.5, 0.5)
	fig, ax = plt.subplots(figsize=(14, 7))

	colors = ['red', 'orange', 'yellow', 'green', 'blue', 'purple',
			  'pink', 'brown', 'black', 'cyan', 'magenta', 'maroon', 'olive',
			  'navy', 'teal', 'lime', 'indigo', 'violet', 'coral']
	linestyles = ['-', '--', '-.', ':', '-', '--', '-.', ':', '-', '--']
	valid_count = min(len(WIND_INSTRUMENT_HEIGHTS), len(ENERGY_FREQ))

	for i in range(valid_count):
		height = WIND_INSTRUMENT_HEIGHTS[i]
		ax.plot(wind_speeds, ENERGY_FREQ[i], label=f'{height}m',
				color=colors[i], linestyle=linestyles[i], linewidth=2, markersize=1)

	ax.set_xlabel('风速 (m/s)', fontsize=12)
	ax.set_ylabel('风能频率 (%)', fontsize=12)
	ax.set_title('测风塔各高度风能频率分布', fontsize=14, fontweight='bold')

	ax.set_xticks(range(0, 26, 2))
	ax.set_xlim(0.5, 25)
	if ENERGY_FREQ.size > 0:
		ax.set_ylim(0, ENERGY_FREQ.max() * 1.1)

	ax.legend(fontsize=10, loc='upper right')
	ax.grid(True, alpha=0.3, linestyle='--')

	return save_fig("fig_energy_frequency")


def plot_weibull_H():
	"""绘制TOWER_HEIGHTm高度Weibull曲线（图2.8-1）"""
	set_chinese_font()
	wind_speeds = np.linspace(0, 20, 100)
	weibull_pdf = (WEIBULL_K[0] / WEIBULL_A[0]) * (wind_speeds / WEIBULL_A[0]) ** (WEIBULL_K[0] - 1) * np.exp(
		-(wind_speeds / WEIBULL_A[0]) ** WEIBULL_K[0])

	fig, ax = plt.subplots(figsize=(10, 6))
	ax.bar(measured_wind, measured_freq, width=0.4, label='实测风速频率', alpha=0.7, color='skyblue')
	ax.plot(wind_speeds, weibull_pdf * 100, 'r-', label=f'Weibull曲线（k={WEIBULL_K[0]}, c={WEIBULL_A[0]}）',
			linewidth=2)

	ax.set_xlabel('风速 (m/s)', fontsize=12)
	ax.set_ylabel('频率 (%)', fontsize=12)
	ax.set_title("测风塔" + TOWER_HEIGHT + "m高度Weibull曲线", fontsize=14, fontweight='bold')
	ax.legend(fontsize=10)
	ax.grid(True, alpha=0.3)
	return save_fig("fig_weibull_H")


def parse_wind_data(data_str):
	"""解析风向/风能数据字符串，返回高度和16扇区数据"""
	parts = list(map(float, data_str.split(',')))
	height = parts[0]  # 第一个值为高度
	sectors = parts[1:17]  # 接下来16个值为扇区数据
	return height, sectors


def plot_wind_energy_roses(wind_freq_data, energy_freq_data, num_heights):
	"""绘制风向玫瑰图和风能玫瑰图（16扇区）"""
	set_chinese_font()
	directions = ["N", "NNE", "NE", "ENE", "E", "ESE", "SE", "SSE",
				  "S", "SSW", "SW", "WSW", "W", "WNW", "NW", "NNW"]
	angles = np.linspace(0, 2 * np.pi, 16, endpoint=False)
	directions += [directions[0]]
	angles = np.concatenate((angles, [angles[0]]))

	fig, axes = plt.subplots(num_heights, 2, figsize=(12, 5 * num_heights),
							 subplot_kw=dict(polar=True))
	if num_heights == 1:
		axes = np.array([axes])

	for row in range(num_heights):
		wind_height, wind_sectors = parse_wind_data(wind_freq_data[row])
		energy_height, energy_sectors = parse_wind_data(energy_freq_data[row])

		wind_sectors = np.concatenate((wind_sectors, [wind_sectors[0]]))
		energy_sectors = np.concatenate((energy_sectors, [energy_sectors[0]]))

		ax_wind = axes[row, 0]
		ax_wind.bar(angles, wind_sectors, width=2 * np.pi / 16,
					bottom=0.0, alpha=0.6, color='skyblue')
		ax_wind.set_xticks(angles[:-1])
		ax_wind.set_xticklabels(directions[:-1], fontsize=8)
		ax_wind.set_title(f'{wind_height}m高度 风向频率玫瑰图', fontsize=12, fontweight='bold')

		ax_energy = axes[row, 1]
		ax_energy.bar(angles, energy_sectors, width=2 * np.pi / 16,
					  bottom=0.0, alpha=0.6, color='salmon')
		ax_energy.set_xticks(angles[:-1])
		ax_energy.set_xticklabels(directions[:-1], fontsize=8)
		ax_energy.set_title(f'{energy_height}m高度 风能频率玫瑰图', fontsize=12, fontweight='bold')

	plt.tight_layout()
	return save_fig("fig_wind_energy_roses")


def set_paragraph_format(paragraph, font_name_cn='宋体', font_name_en='Times New Roman', font_size=12,
						 first_line_indent=2):
	"""设置段落格式：中文宋体，英文Times New Roman，首行缩进"""
	# 设置段落整体字体
	paragraph.style.font.name = font_name_en
	paragraph.style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
	paragraph.style.font.size = Pt(font_size)

	# 设置首行缩进（2个字符）
	paragraph_format = paragraph.paragraph_format
	paragraph_format.first_line_indent = Inches(0.2 * first_line_indent)  # 约等于2个字符
	paragraph_format.space_after = Pt(0)
	paragraph_format.line_spacing = 1.5  # 1.5倍行距

	# 处理段落中的每个run，确保中英文分别使用正确字体
	for run in paragraph.runs:
		run.font.name = font_name_en
		run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
		run.font.size = Pt(font_size)

def insert_centered_image(doc, image_path, width=None):
	"""插入居中显示的图片"""
	# 添加一个新段落用于放置图片
	para = doc.add_paragraph()
	# 设置段落居中对齐
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	# 插入图片
	run = para.add_run()
	run.add_picture(image_path, width=width)
	# 可以根据需要添加图片下方的说明文字
	return para

def add_page_number(section,firstPage):
	"""为指定节添加页码"""
	footer = section.footer
	p = footer.paragraphs[0]
	p.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 页码居中

	# 设置页码格式
	run = p.add_run()
	fldChar = OxmlElement('w:fldChar')
	fldChar.set(qn('w:fldCharType'), 'begin')
	run._element.append(fldChar)

	instrText = OxmlElement('w:instrText')
	instrText.text = 'PAGE'
	run._element.append(instrText)

	fldChar = OxmlElement('w:fldChar')
	fldChar.set(qn('w:fldCharType'), 'end')
	run._element.append(fldChar)

	# 设置页码字体
	run.font.name = 'Times New Roman'
	run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
	run.font.size = Pt(10)

	sectPr = section._sectPr
	pgNumType = OxmlElement('w:pgNumType')
	pgNumType.set(qn('w:start'), firstPage)
	titlePg = OxmlElement('w:titlePg')
	sectPr.append(titlePg)
	sectPr.append(pgNumType)



# -------------------------- 3. Word文档生成函数 --------------------------
# 设置标题样式（修复标题1和标题2字体）
def set_heading_styles(doc):
	# 标题1样式（一级标题）
	heading1 = doc.styles['Heading 1']
	heading1.font.name = '微软雅黑'
	heading1.font.size = Pt(16)
	heading1.font.bold = True
	heading1.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
	# 设置中文字体
	heading1._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

	# 标题2样式（二级标题）
	heading2 = doc.styles['Heading 2']
	heading2.font.name = '微软雅黑'
	heading2.font.size = Pt(14)
	heading2.font.bold = True
	heading2.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
	heading2._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def create_docx():
	"""创建完整的风资源评估报告docx文档"""
	# 1. 初始化文档
	doc = Document()

	set_heading_styles(doc)

	# 设置默认样式：中文宋体，英文Times New Roman
	for style in doc.styles:
		if style.name in ['Normal', 'Heading 1', 'Heading 2', 'Heading 3', 'Table Grid']:
			style.font.name = 'Times New Roman'
			style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
			style.font.size = Pt(12)
			# 标题设置为黑色
			if style.name in ['Heading 1', 'Heading 2']:
				style.font.color.rgb = RGBColor(0, 0, 0)  # 黑色
				style.font.name = '宋体'
				style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
				style.font.bold = True
				if style.name == 'Heading 1':
					style.font.size = Pt(18)  # 一级标题字号加大
				else:
					style.font.size = Pt(14)

	# 封面页设置（单独一节，无页码）
	cover_section = doc.sections[0]
	cover_section.different_first_page_header_footer = True
	cover_section.footer.is_linked_to_previous = False

	# 封面内容
	# 空段落调整标题位置
	for _ in range(8):
		doc.add_paragraph()

	# 主标题
	title = doc.add_heading(LOCATION + TOWER_NAME + '测风塔风资源评估报告', 0)
	title.alignment = WD_ALIGN_PARAGRAPH.CENTER
	for run in title.runs:
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
		run.font.size = Pt(24)
		run.font.bold = True

	# 空段落调整间距
	for _ in range(10):
		doc.add_paragraph()

	# 公司名称
	company = doc.add_paragraph('上海曜疆科技有限公司')
	company.alignment = WD_ALIGN_PARAGRAPH.CENTER
	set_paragraph_format(company, font_size=14)

	# 日期
	dataWord = datetime.now().strftime('%Y/%m')
	date_para = doc.add_paragraph(dataWord)
	date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	set_paragraph_format(date_para, font_size=14)

	# 封面后分页
	doc.add_page_break()

	# 目录页（新节，无页码）
	toc_section = doc.add_section(WD_SECTION.NEW_PAGE)
	toc_section.footer.is_linked_to_previous = False

	# 目录标题
	toc_title = doc.add_heading('目 录', 1)
	toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
	doc.add_paragraph()  # 空行

	# 目录内容（固定文字，实际使用时可在Word中按F9更新页码）
	toc_items = [
		('1 测风塔基本情况', 1),
		('2 测风塔风资源参数', 2),
		('2.1 空气密度', 2),
		('2.2 切变值', 2),
		('2.3 湍流值统计表', 3),
		('2.4 测风塔50年一遇最大风速', 4),
		('2.5 平均风速及风功率密度月变化', 5),
		('2.6 平均风速及风功率密度日变化', 7),
		('2.7 风速及风能频率分布', 8),
		('2.8 Weibull曲线', 9),
		('2.9 风向风能玫瑰图', 10),
		('3 风资源评估结论', 11),
		('3.1 关键风资源参数汇总',11),
		('3.2 评估结论', 11)
	]

	toc_para = doc.add_paragraph()
	for text, page in toc_items:
		# 目录项格式：文字 + 点线 + 页码
		run = toc_para.add_run(text)
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
		# 一级标题目录项加粗加大
		if text.startswith(tuple(f'{i} ' for i in range(1, 10))):
			run.font.bold = True
			run.font.size = Pt(14)

		# 添加点线填充
		tab_stops = toc_para.paragraph_format.tab_stops
		tab_stop = tab_stops.add_tab_stop(Inches(6), WD_ALIGN_PARAGRAPH.RIGHT)
		tab_stop_leader = OxmlElement('w:leader')
		tab_stop_leader.set(qn('w:val'), 'dot')
		tab_stop._element.append(tab_stop_leader)

		toc_para.add_run('\t' + str(page) + '\n')

	set_paragraph_format(toc_para, first_line_indent=0)
	doc.add_page_break()

	# 正文部分（新节，添加页码，从1开始）
	# content_section = doc.add_section(WD_SECTION.NEW_PAGE)
	# content_section.footer.is_linked_to_previous = False
	# add_page_number(content_section)  # 添加页码
	add_page_number(toc_section,'0')

	# -------------------------- 章节1：测风塔基本情况 --------------------------
	para = doc.add_heading('1 测风塔基本情况', 1)
	for run in para.runs:
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

	para = doc.add_paragraph(LOCATION + TOWER_NAME + '测风塔具体配置情况见表1-1。')
	set_paragraph_format(para)

	# 插入表1-1：测风塔配置表
	table1 = doc.add_table(rows=2, cols=7)
	table1.style = 'Table Grid'
	# 表头
	hdr_cells = table1.rows[0].cells
	hdr_cells[0].text = '序号'
	hdr_cells[1].text = '地理坐标'
	hdr_cells[2].text = '海拔高度(m)'
	hdr_cells[3].text = '塔高(m)'
	hdr_cells[4].text = '仪器配置'
	hdr_cells[5].text = '仪器'
	hdr_cells[6].text = '测风时段'
	# 数据行
	row_cells = table1.rows[1].cells
	row_cells[0].text = TOWER_ID
	row_cells[1].text = f'{LATITUDE}\n{LONGITUDE}'
	row_cells[2].text = str(ALTITUDE)
	row_cells[3].text = str(TOWER_HEIGHT)
	row_cells[
		4].text = f'风速仪：{",".join(map(str, WIND_INSTRUMENT_HEIGHTS))}m\n风向仪：{",".join(map(str, WIND_DIR_HEIGHTS))}m\n温度计：{TEMP_PRESS_HEIGHT}m\n气压计：{TEMP_PRESS_HEIGHT}m'
	row_cells[5].text = INSTRUMENT_BRAND
	row_cells[6].text = MEASURE_PERIOD_FULL

	# 设置表格字体（首行首列加粗）
	for row_idx, row in enumerate(table1.rows):
		for col_idx, cell in enumerate(row.cells):
			for para in cell.paragraphs:
				set_paragraph_format(para, first_line_indent=0)
				# 首行或首列加粗
				if row_idx == 0 or col_idx == 0:
					for run in para.runs:
						run.font.bold = True

	doc.add_page_break()

	# -------------------------- 章节2：测风塔风资源参数 --------------------------
	para = doc.add_heading('2 测风塔风资源参数', 1)
	for run in para.runs:
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

	para = doc.add_paragraph(
		f'采用' + LOCATION + TOWER_NAME + f'测风塔{MEASURE_PERIOD_ANALYSIS}的测风数据，进行风资源参数分析，结果如下：')
	set_paragraph_format(para)

	# 2.1 空气密度
	para = doc.add_heading('2.1 空气密度', 2)
	for run in para.runs:
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

	para = doc.add_paragraph(f'用测风塔{TEMP_PRESS_HEIGHT}m高度处实测年温度和气压计算当地的空气密度，计算如下：')
	set_paragraph_format(para)

	# 添加空气密度公式（前面空两格）
	formula_para = doc.add_paragraph()
	formula_para.add_run('ρ = P/(R*T) ，其中：').bold = True
	# 换行并添加首行缩进
	set_paragraph_format(formula_para, first_line_indent=2)  # 设置首行缩进两格

	para = doc.add_paragraph('P—年平均大气压力（Pa）；')
	set_paragraph_format(para)
	para = doc.add_paragraph('R—气体常数（287J/kg·K）；')
	set_paragraph_format(para)
	para = doc.add_paragraph('T—年平均空气开氏温标绝对温度（℃＋273）')
	set_paragraph_format(para)

	# 标红显示空气密度结果
	result_para = doc.add_paragraph('风电场的空气密度')
	result_para.add_run(str(AIR_DENSITY)).font.color.rgb = RGBColor(255, 0, 0)
	result_para.add_run('kg/m³。')
	set_paragraph_format(result_para)

	# 2.2 切变值
	para = doc.add_heading('2.2 切变值', 2)
	for run in para.runs:
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
	para = doc.add_paragraph(
		'在大气边界层中，平均风速随高度发生变化，其变化规律称为风切变；切变指数表示风速在垂直于风向平面内的变化，其大小反映风速随高度增加的快慢，切变的大小取决于测风塔位置的地表粗糙度和大气的稳定程度。测风塔切变值计算结果见下表2.2-1。')
	set_paragraph_format(para)

	# 插入表2.2-1：风切变统计表
	table2 = doc.add_table(rows=len(shear_heights), cols=len(shear_heights))
	table2.style = 'Table Grid'
	# 表头
	hdr_cells = table2.rows[0].cells
	hdr_cells[0].text = '风切变'
	for i, h in enumerate(shear_heights[:-1]):
		hdr_cells[i + 1].text = f'{h}米'
	# 数据行
	for i, h in enumerate(shear_heights[1:]):
		row_cells = table2.rows[i + 1].cells
		row_cells[0].text = f'{h}米'
		for j, val in enumerate(shear_data[i]):
			row_cells[j + 1].text = str(val) if val != '' else ''

	# 设置表格字体（首行首列加粗）
	for row_idx, row in enumerate(table2.rows):
		for col_idx, cell in enumerate(row.cells):
			for para in cell.paragraphs:
				set_paragraph_format(para, first_line_indent=0)
				# 首行或首列加粗
				if row_idx == 0 or col_idx == 0:
					for run in para.runs:
						run.font.bold = True
	para = doc.add_paragraph()
	set_paragraph_format(para)

	# 标红显示综合风切变系数
	para = doc.add_paragraph('经曲线拟合得测风塔综合风切变α=')
	para.add_run(str(SHEAR_COEFFICIENT)).font.color.rgb = RGBColor(255, 0, 0)
	para.add_run('，拟合曲线图见下图。')
	set_paragraph_format(para)

	# 插入风廓线图
	fig1_path = plot_wind_profile()
	insert_centered_image(doc, fig1_path, width=Inches(5.4))
	para = doc.add_paragraph('图2.2-1 测风塔实测风廓线曲线')
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	set_paragraph_format(para, first_line_indent=0)

	# 2.3 湍流值统计表
	para = doc.add_heading('2.3 湍流值统计表', 2)
	for run in para.runs:
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
	para = doc.add_paragraph(
		'湍流强度是描述风速随时间和空间变化的程度，湍流强度越大，表明风速波动越剧烈，对风机的疲劳和寿命影响越大。15m/s风速时各高度湍流强度强度统计如下表2.3-1。')
	set_paragraph_format(para)

	# 插入湍流强度表
	table3 = doc.add_table(rows=2, cols=COLUMNLENGTH + 1)
	table3.style = 'Table Grid'
	hdr_cells = table3.rows[0].cells
	hdr_cells[0].text = '测风高度'
	for i, h in enumerate(WIND_INSTRUMENT_HEIGHTS):
		hdr_cells[i + 1].text = f'{h}米'

	row_cells = table3.rows[1].cells
	row_cells[0].text = '湍流强度'

	for i, val in enumerate(turbulence_data):
		cell = row_cells[i + 1]
		cell.text = ''
		para = cell.add_paragraph()
		run = para.add_run(str(val))
		if i == 0:
			run.font.color.rgb = RGBColor(255, 0, 0)
		set_paragraph_format(para, first_line_indent=0)

	# 设置表格字体（首行首列加粗）
	for row_idx, row in enumerate(table3.rows):
		for col_idx, cell in enumerate(row.cells):
			for para in cell.paragraphs:
				# 首行或首列加粗
				if row_idx == 0 or col_idx == 0:
					for run in para.runs:
						run.font.bold = True

	para = doc.add_paragraph()
	set_paragraph_format(para)

	# 插入TOWER_HEIGHTm高度湍流曲线
	fig2_path = plot_turbulence_H()
	insert_centered_image(doc, fig2_path, width=Inches(5.4))
	para = doc.add_paragraph(f'图2.3-1 测风塔{TOWER_HEIGHT}m高度各风速段湍流强度曲线图')
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	set_paragraph_format(para, first_line_indent=0)

	# 2.4 50年一遇最大风速
	para = doc.add_heading('2.4 测风塔50年一遇最大风速', 2)
	for run in para.runs:
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
	para = doc.add_paragraph('根据独立风暴法计算测风塔50年一遇最大值如下图：')
	set_paragraph_format(para)

	fig3_path = plot_max_wind_50year()
	insert_centered_image(doc, fig3_path, width=Inches(5.4))
	para = doc.add_paragraph('图2.4-1 测风塔50年一遇最大风速')
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	set_paragraph_format(para, first_line_indent=0)

	para = doc.add_paragraph(f'采用独立风暴法推算测风塔{TOWER_HEIGHT}m高度50年一遇最大风速为')
	para.add_run(str(MAX_WIND_50YEAR_H)).font.color.rgb = RGBColor(255, 0, 0)
	para.add_run('m/s。')
	set_paragraph_format(para)

	# 2.5 月平均风速及风功率密度
	para = doc.add_heading('2.5 月平均风速及风功率密度月变化', 2)
	for run in para.runs:
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
	para = doc.add_paragraph('测风塔各高度月平均风速及风功率密度统计如下表2.5-1和图2.5-1、2.5-2所示。')
	set_paragraph_format(para)

	# 插入月平均风速表
	table4 = doc.add_table(rows=13, cols=COLUMNLENGTH + 1)
	table4.style = 'Table Grid'

	# 表头
	hdr_cells = table4.rows[0].cells
	hdr_cells[0].text = '月份'
	for i, h in enumerate(WIND_INSTRUMENT_HEIGHTS):
		hdr_cells[i + 1].text = f'{h}m风速(m/s)'

	# 填充数据
	for month_idx in range(12):
		row = table4.rows[month_idx + 1]
		row.cells[0].text = f'{month_idx + 1}月'
		for i in range(COLUMNLENGTH):
			cell = row.cells[i + 1]
			val = MONTHLY_AVG_WIND[month_idx][i]
			cell.text = ""
			para = cell.add_paragraph()
			run = para.add_run(f'{val:.2f}')
			if i == 0:
				run.font.color.rgb = RGBColor(255, 0, 0)
			set_paragraph_format(para, first_line_indent=0)

	# 插入年平均值行
	avg_row = table4.add_row()
	avg_row.cells[0].text = '全年平均'
	for i in range(COLUMNLENGTH):
		cell = avg_row.cells[i + 1]
		val = MONTHLY_AVG_WIND[:, i].mean()
		cell.text = ""
		para = cell.add_paragraph()
		run = para.add_run(f'{val:.2f}')
		if i == 0:
			run.font.color.rgb = RGBColor(255, 0, 0)
		set_paragraph_format(para, first_line_indent=0)

	# 设置表格字体（首行首列加粗）
	for row_idx, row in enumerate(table4.rows):
		for col_idx, cell in enumerate(row.cells):
			for para in cell.paragraphs:
				# 首行或首列加粗
				if row_idx == 0 or col_idx == 0:
					for run in para.runs:
						run.font.bold = True

	para = doc.add_paragraph()
	set_paragraph_format(para)

	# 插入月平均风速图
	fig5_path = plot_monthly_wind()
	insert_centered_image(doc, fig5_path, width=Inches(5.4))
	para = doc.add_paragraph('图2.5-1 测风塔各高度月平均风速')
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	set_paragraph_format(para, first_line_indent=0)

	# 插入月平均风功率密度表
	table5 = doc.add_table(rows=13, cols=COLUMNLENGTH + 1)
	table5.style = 'Table Grid'

	# 表头
	hdr_cells = table5.rows[0].cells
	hdr_cells[0].text = '月份'
	for i, h in enumerate(WIND_INSTRUMENT_HEIGHTS):
		hdr_cells[i + 1].text = f'{h}m功率密度(W/m²)'

	# 填充数据
	for month_idx in range(12):
		row = table5.rows[month_idx + 1]
		row.cells[0].text = f'{month_idx + 1}月'
		for i in range(COLUMNLENGTH):
			cell = row.cells[i + 1]
			val = MONTHLY_AVG_POWER[month_idx][i]
			cell.text = ""
			para = cell.add_paragraph()
			run = para.add_run(f'{val:.0f}')
			if i == 0:
				run.font.color.rgb = RGBColor(255, 0, 0)
			set_paragraph_format(para, first_line_indent=0)

	# 插入年平均值行
	avg_row = table5.add_row()
	avg_row.cells[0].text = '全年平均'
	for i in range(COLUMNLENGTH):
		cell = avg_row.cells[i + 1]
		val = MONTHLY_AVG_POWER[:, i].mean()
		cell.text = ""
		para = cell.add_paragraph()
		run = para.add_run(f'{val:.0f}')
		if i == 0:
			run.font.color.rgb = RGBColor(255, 0, 0)
		set_paragraph_format(para, first_line_indent=0)

	# 设置表格字体（首行首列加粗）
	for row_idx, row in enumerate(table5.rows):
		for col_idx, cell in enumerate(row.cells):
			for para in cell.paragraphs:
				# 首行或首列加粗
				if row_idx == 0 or col_idx == 0:
					for run in para.runs:
						run.font.bold = True

	para = doc.add_paragraph()
	set_paragraph_format(para)

	# 插入月平均风功率密度图
	fig6_path = plot_monthly_power()
	insert_centered_image(doc, fig6_path, width=Inches(5.4))
	para = doc.add_paragraph('图2.5-2 测风塔各高度月平均风功率密度')
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	set_paragraph_format(para, first_line_indent=0)

	# 2.6 平均风速及风功率密度日变化
	para = doc.add_heading('2.6 平均风速及风功率密度日变化', 2)
	for run in para.runs:
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
	para = doc.add_paragraph('测风塔各测风高度风速和风功率密度日变化见图2.6-1和2.6-2。')
	set_paragraph_format(para)

	# 插入日平均风速图
	fig6_1_path = plot_hourly_wind()
	insert_centered_image(doc, fig6_1_path, width=Inches(5.4))
	para = doc.add_paragraph('图2.6-1 测风塔各高度平均风速日变化')
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	set_paragraph_format(para, first_line_indent=0)

	# 插入日平均风功率密度图
	fig6_2_path = plot_hourly_power()
	insert_centered_image(doc, fig6_2_path, width=Inches(5.4))
	para = doc.add_paragraph('图2.6-2 测风塔各高度平均风功率密度日变化')
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	set_paragraph_format(para, first_line_indent=0)

	# 2.7 风速及风能频率分布
	para = doc.add_heading('2.7 风速及风能频率分布', 2)
	for run in para.runs:
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
	para = doc.add_paragraph('测风塔各测风高度风速和风能频率分布见下表和下图。')
	set_paragraph_format(para)

	# 插入风速频率分布图
	fig7_1_path = plot_wind_frequency()
	insert_centered_image(doc, fig7_1_path, width=Inches(5.4))
	para = doc.add_paragraph('图2.7-1 测风塔各高度风速频率分布')
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	set_paragraph_format(para, first_line_indent=0)

	# 插入风能频率分布图
	fig7_2_path = plot_energy_frequency()
	insert_centered_image(doc, fig7_2_path, width=Inches(5.4))
	para = doc.add_paragraph('图2.7-2 测风塔各高度风能频率分布')
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	set_paragraph_format(para, first_line_indent=0)

	# 2.8 Weibull曲线
	para = doc.add_heading('2.8 Weibull曲线', 2)
	for run in para.runs:
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
	para = doc.add_paragraph('各高层weibull参数如表2.8-1所示。')
	set_paragraph_format(para)

	# 插入Weibull参数表
	table6 = doc.add_table(rows=2, cols=COLUMNLENGTH + 1)
	table6.style = 'Table Grid'
	hdr_cells = table6.rows[0].cells
	hdr_cells[0].text = '高度'
	for i, h in enumerate(WIND_INSTRUMENT_HEIGHTS):
		hdr_cells[i + 1].text = f'{h}米'

	# 处理k（形状参数）行
	k_row = table6.rows[1]
	k_row.cells[0].text = 'k（形状参数）'
	for i, k in enumerate(WEIBULL_K):
		k_row.cells[i + 1].text = str(k)
		for para in k_row.cells[i + 1].paragraphs:
			set_paragraph_format(para, first_line_indent=0)

	# 添加A参数行
	a_row = table6.add_row()
	a_row.cells[0].text = 'c（尺度参数）'
	for i, a in enumerate(WEIBULL_A):
		a_row.cells[i + 1].text = str(a)
		for para in a_row.cells[i + 1].paragraphs:
			set_paragraph_format(para, first_line_indent=0)

	for row_idx, row in enumerate(table6.rows):
		for col_idx, cell in enumerate(row.cells):
			for para in cell.paragraphs:
				# 首行或首列加粗
				if row_idx == 0 or col_idx == 0:
					for run in para.runs:
						run.font.bold = True

	para = doc.add_paragraph()
	set_paragraph_format(para)

	# 插入TOWER_HEIGHTm高度Weibull曲线
	fig7_path = plot_weibull_H()
	insert_centered_image(doc, fig7_path, width=Inches(5.4))
	para = doc.add_paragraph(f'图2.8-1 测风塔{TOWER_HEIGHT}m高度Weibull曲线')
	para.alignment = WD_ALIGN_PARAGRAPH.CENTER
	set_paragraph_format(para, first_line_indent=0)

	# 2.9 风向风能玫瑰图
	para = doc.add_heading('2.9 风向风能玫瑰图', 2)
	for run in para.runs:
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
	para = doc.add_paragraph('测风塔各高度风向频率与风能频率玫瑰图如下：')
	set_paragraph_format(para)
	# 插入玫瑰图（指定显示前2层高度）
	fig8_path = plot_wind_energy_roses(WIND_FREQ_STRINGS, ENERGY_FREQ_STRINGS, num_heights=wd_heights)
	insert_centered_image(doc, fig8_path, width=Inches(5.2))
	doc.add_paragraph('图2.9-1 各高度风向风能玫瑰图').alignment = WD_ALIGN_PARAGRAPH.CENTER

	doc.add_page_break()
	# -------------------------- 章节3：风资源评估结论 --------------------------
	para = doc.add_heading('3 风资源评估结论', 1)
	for run in para.runs:
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
	para = doc.add_paragraph(
		'通过对' + LOCATION + TOWER_NAME + '测风塔风资源数据的系统分析，得出以下主要结论：'.format(
			TOWER_ID=TOWER_ID))
	set_paragraph_format(para)

	# 3.1 关键风资源参数汇总
	para = doc.add_heading('3.1 关键风资源参数汇总', 2)
	for run in para.runs:
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
	# 插入关键参数汇总表（标红TOWER_HEIGHTm高度核心数据）
	summary_table = doc.add_table(rows=7, cols=3)
	summary_table.style = 'Table Grid'

	# 表头
	hdr_cells = summary_table.rows[0].cells
	hdr_cells[0].text = '参数名称'
	hdr_cells[1].text = TOWER_HEIGHT + 'm高度数值'
	hdr_cells[2].text = '单位'

	# 参数数据
	params = [
		('年平均风速', AVG_WIND_SPEED_H, 'm/s'),
		('空气密度', AIR_DENSITY, 'kg/m³'),
		('综合风切变系数', SHEAR_COEFFICIENT, ''),
		('15m/s时平均湍流强度', TURBULENCE_15MPS_H, ''),
		('15m/s时代表湍流强度', TURBULENCE_CHAR_15MPS_H, ''),
		('50年一遇最大风速', MAX_WIND_50YEAR_H, 'm/s')
	]

	for row_idx, (name, value, unit) in enumerate(params):
		row = summary_table.rows[row_idx + 1]

		# 填充参数名称
		name_cell = row.cells[0]
		name_cell.text = name

		# 填充TOWER_HEIGHTm高度数值（标红）
		value_cell = row.cells[1]
		value_cell.text = ""  # 清空单元格
		para = value_cell.add_paragraph()
		run = para.add_run(f'{value:.3f}' if isinstance(value, float) else str(value))
		run.font.color.rgb = RGBColor(255, 0, 0)  # 标红显示

		# 填充单位
		unit_cell = row.cells[2]
		unit_cell.text = unit

	for row_idx, row in enumerate(summary_table.rows):
		for col_idx, cell in enumerate(row.cells):
			for para in cell.paragraphs:
				# 首行或首列加粗
				if row_idx == 0 or col_idx == 0:
					for run in para.runs:
						run.font.bold = True

	# 3.2 风资源评估结论
	para = doc.add_heading('3.2 评估结论', 2)
	for run in para.runs:
		run.font.name = 'Times New Roman'
		run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
	conclusion_para1 = doc.add_paragraph()
	conclusion_para1.add_run('1、 风速特性：').bold = True
	conclusion_para1.add_run(
		f'测风塔{TOWER_HEIGHT}m高度年平均风速为{AVG_WIND_SPEED_H}m/s，属于中等风资源区域。风速月变化呈现明显季节性特征。')

	conclusion_para2 = doc.add_paragraph()
	conclusion_para2.add_run('2、 风功率潜力：').bold = True
	conclusion_para2.add_run(
		f'{TOWER_HEIGHT}m高度年平均风功率密度约为{MONTHLY_AVG_POWER[:, 0].mean():.0f}W/m²，根据风资源评估标准，该区域具备一定的风能开发价值。')

	conclusion_para3 = doc.add_paragraph()
	conclusion_para3.add_run('3、 湍流特性：').bold = True
	conclusion_para3.add_run(
		f'{TOWER_HEIGHT}m高度15m/s风速时平均湍流强度为{TURBULENCE_15MPS_H}，代表湍流强度为{TURBULENCE_CHAR_15MPS_H}，均在正常范围内，对风机运行影响可控。')

	conclusion_para4 = doc.add_paragraph()
	conclusion_para4.add_run('4、 极端风速：').bold = True
	conclusion_para4.add_run(
		f'{TOWER_HEIGHT}m高度50年一遇最大风速为{MAX_WIND_50YEAR_H:.2f}m/s，风机选型时需满足该极端风速条件。')

	conclusion_para5 = doc.add_paragraph()
	conclusion_para5.add_run('5、 综合评价：').bold = True
	conclusion_para5.add_run(
		'该测风塔所在区域风资源条件总体良好，主风向稳定，湍流强度适中，具备风力发电开发的基本条件，建议进行后续可行性研究。')

	# 保存文档
	doc.save(f'{LOCATION}{TOWER_NAME}测风塔风资源评估报告.docx')
	print("报告生成完成！")


if __name__ == "__main__":
	create_docx()